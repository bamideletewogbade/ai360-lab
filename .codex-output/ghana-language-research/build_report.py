from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "report-source.md"
TEMPLATE = Path(r"C:\Users\HP\.codex\plugins\cache\openai-curated-remote\openai-templates\0.1.1\skills\artifact-template-system-design\assets\reference.docx")
OUTPUT_DIR = Path(r"C:\Users\HP\Desktop\AIC\ai360-lab\artifacts")
OUTPUT = OUTPUT_DIR / "AI360-Ghanaian-Language-Voice-Research.docx"

NAVY = "17324D"
BLUE = "52718E"
PALE_BLUE = "DCEAF5"
PALEST_BLUE = "F3F7FA"
MID_GRAY = "D5DEE7"
TEXT = "17202A"
MUTED = "5A6A78"
WHITE = "FFFFFF"
AMBER = "F3E4B3"
GREEN = "DDEEDC"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color=MID_GRAY, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_cell_text(cell, text: str, *, bold=False, color=TEXT, size=8.5) -> None:
    text = text.replace("`", "")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def add_hyperlink(paragraph, text: str, url: str, color=BLUE, underline=True):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("AI360  •  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 6),
        ("Heading 1", 18, NAVY, 14, 7),
        ("Heading 2", 13.5, BLUE, 11, 5),
        ("Heading 3", 11, NAVY, 9, 3),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Title" or True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        try:
            style = styles[list_name]
        except KeyError:
            style = styles.add_style(list_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = normal
        style.font.name = "Aptos"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.16)


def add_inline(paragraph, text: str) -> None:
    # URLs become normal hyperlinks. Markdown emphasis and code become styled runs.
    url_re = re.compile(r"https?://[^\s)]+")
    token_re = re.compile(r"(\*\*.*?\*\*|`.*?`|https?://[^\s)]+)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(NAVY)
        elif url_re.fullmatch(token):
            add_hyperlink(paragraph, "source", token)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_source_item(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.add_run("• ")
    urls = re.findall(r"https?://\S+", text)
    label = re.split(r"https?://", text, maxsplit=1)[0].rstrip(" :")
    p.add_run(label)
    if urls:
        p.add_run(": ")
        for index, url in enumerate(urls, 1):
            url = url.rstrip(".,")
            if index > 1:
                p.add_run("; ")
            add_hyperlink(p, "link" if len(urls) == 1 else f"link {index}", url)


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for i, row in enumerate(rows):
        for j in range(cols):
            text = row[j] if j < len(row) else ""
            set_cell_text(table.cell(i, j), text, bold=(i == 0), color=WHITE if i == 0 else TEXT, size=8.2)
            shade(table.cell(i, j), NAVY if i == 0 else (PALEST_BLUE if i % 2 == 0 else WHITE))
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_architecture_flow(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("VOICE → ROUTE → TRANSCRIBE → INTERPRET → RESPOND → LEARN")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    labels = [
        ("1  Capture", "Automatic after stop"),
        ("2  Identify", "Language • dialect • switches"),
        ("3  Recognize", "Per-language ASR policy"),
        ("4  Understand", "Entities • meaning • uncertainty"),
        ("5  Answer", "Direct or pivot realization"),
        ("6  Improve", "Consent-safe corrections"),
    ]
    table = document.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="B9C9D8")
    for idx, (title, subtitle) in enumerate(labels):
        cell = table.cell(idx // 3, idx % 3)
        shade(cell, PALE_BLUE if idx % 2 == 0 else PALEST_BLUE)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(title)
        r1.bold = True
        r1.font.color.rgb = RGBColor.from_string(NAVY)
        r1.font.size = Pt(10)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor.from_string(MUTED)
        set_cell_margins(cell, top=140, start=100, bottom=140, end=100)


def add_cover(document: Document) -> None:
    for _ in range(4):
        document.add_paragraph()
    eyebrow = document.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(8)
    run = eyebrow.add_run("AI360 LANGUAGE INTELLIGENCE")
    run.font.name = "Aptos"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    run.bold = False

    title = document.add_paragraph(style="Title")
    title.paragraph_format.space_after = Pt(7)
    r = title.add_run("Improving Ghanaian Language\nVoice and Translation")
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(4)
    r = subtitle.add_run("Research findings, target architecture, and 90-day implementation plan")
    r.font.name = "Aptos"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    for _ in range(5):
        document.add_paragraph()

    meta = document.add_table(rows=1, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = True
    values = [
        ("STATUS", "Proposed"),
        ("OWNER", "AI360 product + engineering"),
        ("LAST UPDATED", "4 September 2026"),
    ]
    for idx, (label, value) in enumerate(values):
        cell = meta.cell(0, idx)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(3)
        rr = p1.add_run(label)
        rr.bold = True
        rr.font.size = Pt(8)
        rr.font.color.rgb = RGBColor.from_string(BLUE)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        rr = p2.add_run(value)
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor.from_string(TEXT)
        set_cell_margins(cell, 30, 0, 30, 80)
    set_table_borders(meta, color=WHITE, size="0")

    document.add_paragraph()
    scope = document.add_table(rows=3, cols=2)
    scope.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(scope, color=WHITE, size="0")
    cover_rows = [
        ("Authors", "Codex research synthesis for AI360"),
        ("Reviewers", "Ghanaian native speakers, product, engineering, privacy"),
        ("Scope", "Speech recognition, interpretation, translation, response generation, and rollout"),
    ]
    for i, (label, value) in enumerate(cover_rows):
        set_cell_text(scope.cell(i, 0), label, bold=True, size=8.5)
        set_cell_text(scope.cell(i, 1), value, size=8.5)
        shade(scope.cell(i, 0), "C9DEEF")
        shade(scope.cell(i, 1), PALEST_BLUE)
    document.add_page_break()


def add_executive_panel(document: Document) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="9FB6CA")
    cell = table.cell(0, 0)
    shade(cell, PALE_BLUE)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run("RECOMMENDATION")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(
        "Approve a language-aware, evaluation-first voice pipeline. Fix routing and automatic transcription first; benchmark Ghana-specific ASR and direct-vs-pivot language generation before choosing or training production models."
    )
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor.from_string(NAVY)
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)


def build() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(TEMPLATE, OUTPUT)
    document = Document(OUTPUT)
    clear_body(document)
    configure_styles(document)

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.3)
    document.settings.odd_and_even_pages_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "AI360  /  Ghanaian Language Voice System Design"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)
    fp = section.footer.paragraphs[0]
    fp.text = ""
    add_page_number(fp)
    even_header = section.even_page_header
    ehp = even_header.paragraphs[0]
    ehp.text = "AI360  /  Ghanaian Language Voice System Design"
    ehp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in ehp.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)
    efp = section.even_page_footer.paragraphs[0]
    efp.text = ""
    add_page_number(efp)

    add_cover(document)
    add_executive_panel(document)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    in_table = False
    table_rows: list[list[str]] = []
    sources_mode = False
    skip_cover_meta = True

    def flush_table():
        nonlocal table_rows, in_table
        if table_rows:
            add_markdown_table(document, table_rows)
        table_rows = []
        in_table = False

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("# "):
            continue
        if skip_cover_meta and (not line or line.startswith("**Status:") or line.startswith("**Owner:") or line.startswith("**Last updated:") or line.startswith("**Scope:")):
            continue
        if line.startswith("## Executive decision"):
            skip_cover_meta = False
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                continue
            in_table = True
            table_rows.append(cells)
            continue
        if in_table:
            flush_table()
        if not line:
            continue
        if line.startswith("## "):
            heading = line[3:]
            sources_mode = heading == "Sources"
            if heading == "Proposed architecture":
                document.add_page_break()
            document.add_heading(heading, level=1)
            if heading == "Proposed architecture":
                add_architecture_flow(document)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
            continue
        if line.startswith("- "):
            item = line[2:]
            if sources_mode:
                add_source_item(document, item)
            else:
                p = document.add_paragraph(style="List Bullet")
                p.add_run("• ")
                add_inline(p, item)
            continue
        if re.match(r"^\d+\. ", line):
            number, item = line.split(". ", 1)
            p = document.add_paragraph(style="List Number")
            p.add_run(f"{number}. ")
            add_inline(p, item)
            continue
        p = document.add_paragraph()
        add_inline(p, line)

    flush_table()

    # A compact closing signal for reviewers.
    document.add_paragraph()
    table = document.add_table(rows=1, cols=1)
    set_table_borders(table, color="A9C5A7")
    shade(table.cell(0, 0), GREEN)
    set_cell_text(
        table.cell(0, 0),
        "Approval requested: proceed with the 0–30 day measurement and routing foundation before committing to a production model or fine-tuning programme.",
        bold=True,
        color=NAVY,
        size=9.5,
    )
    endp = document.add_paragraph()
    endp.paragraph_format.space_after = Pt(0)
    endp.paragraph_format.line_spacing = Pt(1)
    endrun = endp.add_run(" ")
    endrun.font.size = Pt(1)

    # Avoid isolated headings and table row splits.
    for p in document.paragraphs:
        if p.style and p.style.name.startswith("Heading"):
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
    for table in document.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)

    document.core_properties.title = "Improving Ghanaian Language Voice and Translation in AI360"
    document.core_properties.subject = "Research, architecture, evaluation, and 90-day implementation plan"
    document.core_properties.author = "AI360 / Codex"
    document.core_properties.keywords = "AI360, Ghana, Twi, Ewe, Ga, Ghanaian Pidgin, speech recognition, translation"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
